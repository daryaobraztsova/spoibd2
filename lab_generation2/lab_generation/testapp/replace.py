import docx
from docx2python.utilities import replace_docx_text
import os

keyword_dict = {
    "[num]": "",
    "[group]": "",
    "[username]": "",
    "[username1]": "",
    "[mission]": ""
}

input_file = "akadem4.docx"
output_file = "out.docx"

def checkIfExists():
    if not os.path.isfile(output_file):
        document = docx.Document(input_file)
        document.save(output_file)

def replace():
    for i in keyword_dict:
        doc = docx.Document()
        doc.save("tmp.docx")
        tmp_file = "tmp.docx"
        replace_docx_text(
            output_file,
            tmp_file,
            (i, keyword_dict[i]),  # replace Apples with Bananas
            html=True,
        )
        doc = docx.Document("tmp.docx")
        doc.save(output_file)
    os.remove(tmp_file)
