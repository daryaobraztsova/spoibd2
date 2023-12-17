import docx
import os

from .replace import *

output_file2 = "out2.docx"

def newOut():
    doc = docx.Document(output_file)
    doc.save(output_file2)

def writeToDocx(text, isBold):
    doc = docx.Document(output_file2)
    p = doc.add_paragraph()
    runner = p.add_run(text)
    runner.bold = isBold
    doc.save(output_file2)